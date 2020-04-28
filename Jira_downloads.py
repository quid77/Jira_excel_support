import collections
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException, StaleElementReferenceException, TimeoutException
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
import pathlib
import win32com
import time
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import os
import shutil
from docx.shared import Inches, Pt
import glob
import re
import pandas
import win32com.client as win32
from win32com.client import constants
import docx, docxpy
import unittest

user_login = ""
user_password = ""


script_path = pathlib.Path(__file__).parent.absolute()
download_path = str(script_path) + "\\Downloads"
titles_and_tasks = collections.defaultdict(list)


def set_up_excel_data():
    file_handler = pandas.read_excel(str(script_path) + "\\Book.xlsx")

    all_labels = [label for label in file_handler.columns]
    for label in all_labels:
        data = pandas.DataFrame(file_handler, columns=[label])
        column = data.to_string(index=False).split('\n')
        stripped_list = [elem.strip().split('\\n') for elem in column]

        for cell in stripped_list:
            issue_link = [link for link in cell if "https" in link]
            if issue_link:
                titles_and_tasks[label].append(issue_link)


class JiraTestsDownload(unittest.TestCase):

    @classmethod
    def setUpClass(self):  # setUpClass runs once for ALL tests
        options = webdriver.ChromeOptions()
        preferences = {"download.default_directory": download_path, "download.prompt_for_download": "false",
                       "safebrowsing.enabled": "false", 'profile.default_content_setting_values.automatic_downloads': 1}
        options.add_experimental_option("prefs", preferences)
        self.driver = webdriver.Chrome(options=options)

    def test_1_login_to_app(self):
        driver = self.driver
        driver.implicitly_wait(10)
        driver.get("https://jira.softserveinc.com/projects/DP/issues/DP-707?filter=allopenissues")
        login_element = driver.find_element_by_id("login-form-username")
        password_element = driver.find_element_by_id("login-form-password")
        login_element.send_keys(user_login, Keys.TAB)
        password_element.click()
        password_element.send_keys(user_password, Keys.ENTER)

    def test_2_download_tests(self):
        driver = self.driver
        wait = WebDriverWait(driver, 1)
        for label, label_links in titles_and_tasks.items():
            for link in label_links:
                driver.get(link[0])
                for x in range(0, 10):
                    try:
                        export = wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@id='viewissue-export']")))
                        driver.execute_script("arguments[0].click();", export)
                        word = wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@id='jira.issueviews:issue-word']")))
                        driver.execute_script("arguments[0].click();", word)
                        break
                    except (StaleElementReferenceException, TimeoutException):
                        time.sleep(0.2)
                        continue

    @classmethod
    def tearDownClass(self):
        time.sleep(3)
        self.driver.close()
        # Replace files if newer version of that file was downloaded
        files = glob.glob(download_path + "\\DP*")
        for file in files:
            if file[:-4].endswith(")"):
                shutil.move(file, file[:-8] + ".doc")


def create_dir_hierarchy():
    if not os.path.exists(download_path + "\\DocFiles"):
        os.makedirs(download_path + "\\DocFiles")
    if not os.path.exists(download_path + "\\DocxFiles"):
        os.makedirs(download_path + "\\DocxFiles")
    if not os.path.exists(download_path + "\\TestTemplates"):
        os.makedirs(download_path + "\\TestTemplates")
    if not os.path.exists(download_path + "\\Directories"):
        os.makedirs(download_path + "\\Directories")


def move_doc_files():
    # Create list of paths to .doc files
    for filename in os.listdir(download_path):
        if filename.endswith(".doc") and filename.startswith("DP"):
            if os.path.exists(download_path + "\\DocFiles\\" + filename):
                os.remove(download_path + "\\DocFiles\\" + filename)
            shutil.move(download_path + "\\" + filename, download_path + "\\DocFiles\\" + filename)


# this function isn't standalone, use save_to_docx instead
def save_as_docx(name):
    # Opening MS Word
    word = win32.gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(download_path + "\\DocFiles\\" + name)
    doc.Activate()

    # Rename path with .docx
    new_file_abs = os.path.abspath(download_path + "\\DocxFiles\\" + name)
    new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

    # Save and Close
    word.ActiveDocument.SaveAs(new_file_abs, FileFormat=constants.wdFormatXMLDocument)
    doc.Close(False)


def save_to_docx():
    file_names = os.listdir(download_path + "\\DocFiles")
    for name in file_names:
        if os.path.exists((download_path + "\\DocxFiles\\" + name)[:-3] + "docx"):
            os.remove((download_path + "\\DocxFiles\\" + name)[:-3] + "docx")
        save_as_docx(name)


def read_docx_files():
    files = glob.glob(download_path + "\\DocxFiles\\DP*")
    os.chdir(download_path)
    for file in files:
        if os.path.exists(download_path + "\\TestTemplates\\" + os.path.basename(file)):
            os.remove(download_path + "\\TestTemplates\\" + os.path.basename(file))
        docx_handler = docx.Document(file)
        docx_tables = docx_handler.tables
        docx_hyperlink_handler = docxpy.DOCReader(file)
        docx_hyperlink_handler.process()
        hyperlinks = docx_hyperlink_handler.data['links']
        test_scenario_hyperlink_text = str(hyperlinks[0][0])[2:-1]
        jira_id = docx_tables[0].rows[0].cells[0].text
        jira_test_id = jira_id.split()[0]

        jira_test_desc = ""
        for i in range(len(docx_tables)):
            if "Description" in docx_tables[i].rows[0].cells[0].text:
                jira_test_desc = docx_tables[i+1].rows[0].cells[0].text.strip()
                break

        zephyr_tests = None
        for x in range(len(docx_tables[2].rows)):
            if "Zephyr" in docx_tables[2].rows[x].cells[0].text:
                zephyr_tests = docx_tables[2].rows[x].cells[1]
                break
        if not zephyr_tests:
            # Move test without steps to separate dir
            if not os.path.exists(download_path + "\\DocxFiles\\TestsWithoutSteps"):
                os.makedirs(download_path + "\\DocxFiles\\TestsWithoutSteps")
            shutil.move(download_path + "\\DocxFiles\\" + os.path.basename(file),
                        download_path + "\\DocxFiles\\TestsWithoutSteps\\" + os.path.basename(file))
            continue
        zephyr_tests_table = zephyr_tests.tables

        zephyr_rows = zephyr_tests_table[0].rows  # get row id's
        zephyr_rows = zephyr_rows[1:]  # remove first cell from all rows (e.g. "Test Step", "Test Data", etc.)

        #  Test Steps
        list_of_test_steps = []
        for row in zephyr_rows:
            list_of_test_steps.append(row.cells[1].text)

        #  Test Conditions
        list_of_test_conditions = []
        for row in zephyr_rows:
            list_of_test_conditions.append(row.cells[2].text)

        #  Expected results
        list_of_exptected_results = []
        for row in zephyr_rows:
            list_of_exptected_results.append(row.cells[3].text)

        number_of_teststeps = zephyr_tests_table[0].rows[-1].cells[0].text
        file_save_path = download_path + "\\TestTemplates\\" + os.path.basename(file)

        final_docx_template = docx.Document(str(script_path) + "\\SampleTestScripts1.docx")
        final_docx_table = final_docx_template.tables
        font = final_docx_template.styles['Normal'].font
        font.name = 'Calibri'
        paragraph = final_docx_template.styles['Normal'].paragraph_format
        paragraph.space_after = Pt(3)
        paragraph.left_indent = Pt(0)
        heading1 = final_docx_template.styles['Heading 1'].paragraph_format
        heading1.space_before = Pt(0)
        for x in range(1, int(number_of_teststeps)):
            final_docx_table[0].add_row()

        os.chdir(download_path)
        test_id = final_docx_table[0].rows[0].cells[2].paragraphs[0]
        test_id.add_run(jira_test_id)
        test_scenario = final_docx_table[0].rows[1].cells[2].paragraphs[0]
        test_scenario.add_run(test_scenario_hyperlink_text)
        final_docx_table[0].rows[1].cells[2].paragraphs[0].paragraph_format.left_indent = Pt(0)
        test_description = final_docx_table[0].rows[2].cells[2].paragraphs[0]
        test_description.add_run(jira_test_desc)

        steps_only_table = final_docx_table[0].rows[4:]

        # steps_only_table[0].cells[0].paragraphs[0].style = rws_template.styles['Normal']
        for x in range(0, int(number_of_teststeps)):
            steps_only_table[x].cells[0].paragraphs[0].paragraph_format.left_indent = Pt(12)

        for x in range(0, int(number_of_teststeps)):
            steps_only_table[x].cells[0].paragraphs[0].add_run(str(x + 1) + ".")
        for x in range(0, int(number_of_teststeps)):
            steps_only_table[x].cells[1].paragraphs[0].add_run(list_of_test_steps[x])
        for x in range(0, int(number_of_teststeps)):
            steps_only_table[x].cells[2].paragraphs[0].add_run(list_of_test_conditions[x])
        for x in range(0, int(number_of_teststeps)):
            steps_only_table[x].cells[3].paragraphs[0].add_run(list_of_exptected_results[x])

        # final_docx_template.add_page_break()
        final_docx_template.save(file_save_path)


def move_files_to_epics():
    os.chdir(download_path + "\\TestTemplates")

    for label, label_links in titles_and_tasks.items():
        label_name = label
        if not os.path.exists(download_path + "\\Directories\\" + label_name):
            os.makedirs(download_path + "\\Directories\\" + label_name)

        for link in label_links:
            filename = link[0].split('browse/')[1]
            if os.path.exists(download_path + "\\TestTemplates\\" + filename + ".docx"):  # unnecessary if
                shutil.move(download_path + "\\TestTemplates\\" + filename + ".docx",
                            download_path + "\\Directories\\" + label_name + "\\" + filename + ".docx")


def merge_files_in_epics():
    label_dirs = glob.glob(download_path + "\\Directories\\*")
    for label_dir in label_dirs:
        if os.path.basename(label_dir) not in titles_and_tasks.keys() or not os.path.isdir(label_dir):
            break
        os.chdir(label_dir)
        files = glob.glob(label_dir + "\\DP*")
        epic_docx = docx.Document(str(script_path) + "\\EpicTemplate.docx")
        epic_docx.add_page_break()

        font = epic_docx.styles['Normal'].font
        font.name = 'Calibri'
        paragraph = epic_docx.styles['Normal'].paragraph_format
        paragraph.space_after = Pt(3)
        paragraph.left_indent = Pt(0)
        heading1 = epic_docx.styles['Heading 1'].paragraph_format
        heading1.space_before = Pt(0)

        for index, file_path in enumerate(files):
            docx_handler = docx.Document(file_path)

            if index < len(files)-1:
                docx_handler.add_page_break()
            for index, element in enumerate(docx_handler.element.body):
                epic_docx.element.body.append(element)
                if index == 4:
                    break
        epic_docx.save(label_dir + "\\" + os.path.basename(label_dir) + ".docx")
        time.sleep(2)
        word = win32com.client.gencache.EnsureDispatch("Word.Application")
        doc = word.Documents.Open(label_dir + "\\" + os.path.basename(label_dir) + ".docx")
        doc.TablesOfContents(1).Update()
        doc.Close(SaveChanges=True)
        word.Quit()


if __name__ == "__main__":

    create_dir_hierarchy()
    set_up_excel_data()
    unittest.main(exit=False)  # Dont stop the program after test execution (it would skip below functions)
    move_doc_files()
    save_to_docx()
    read_docx_files()
    move_files_to_epics()
    merge_files_in_epics()
